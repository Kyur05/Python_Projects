import docx

def inject_data_to_word(template_path, output_path, data_dictionary):
    print(f"\n>> Opening template: {template_path}...")
    
    # 1. Load the Word Document
    doc = docx.Document(template_path)
    
    # 2. Hunt through standard paragraphs
    for para in doc.paragraphs:
        for tag, real_value in data_dictionary.items():
            if tag in para.text:
                # Replace the tag with the actual engineering data
                para.text = para.text.replace(tag, str(real_value))
                
    # 3. Hunt through tables (Word treats tables separately from paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for tag, real_value in data_dictionary.items():
                        if tag in para.text:
                            para.text = para.text.replace(tag, str(real_value))
                            
    # 4. Save the freshly written report
    doc.save(output_path)
    print(f"[SUCCESS] Automated Report Generated: {output_path}")

if __name__ == "__main__":
    # This is where we will eventually feed the data from your Civil 3D CSVs.
    # For now, we use test data to prove the bridge works.
    
    project_data = {
        "<<PROJECT_NAME>>": "N2 Section 24 Interchange Upgrade",
        "<<TOTAL_CUT>>": "145,230",
        "<<TOTAL_FILL>>": "89,400"
    }
    
    inject_data_to_word("Report_Template.docx", "N2_Draft_Report_Rev0.docx", project_data)